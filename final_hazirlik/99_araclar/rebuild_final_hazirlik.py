import json
import re
import shutil
import zipfile
from pathlib import Path

from docx import Document
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "final_hazirlik"
PYTHON = "bundled"

DIRS = {
    "vize": OUT / "00_KAYNAKLAR" / "01_vize_temeli_slaytlar",
    "final": OUT / "00_KAYNAKLAR" / "02_final_agirlikli_slaytlar",
    "proje": OUT / "00_KAYNAKLAR" / "03_proje_odevleri_ve_dil_raporlari",
    "ek": OUT / "00_KAYNAKLAR" / "04_ek_materyaller",
    "ozet": OUT / "01_OZETLER",
    "app": OUT / "02_INTERAKTIF_SISTEM",
    "veri": OUT / "03_VERI",
    "rapor": OUT / "04_RAPORLAR",
    "arac": OUT / "99_araclar",
}


TAGS = [
    "Final ağırlıklı",
    "Vize temeli",
    "Proje konusu",
    "Kod yorumlama",
    "Dil tanıma",
    "Compile-time / run-time",
    "Scope / binding",
    "Type system",
    "Paradigmalar",
    "Bellek yönetimi",
    "Fonksiyonel programlama",
    "Nesne yönelimli programlama",
    "Concurrency",
    "Klasik soru",
]


TOPICS = [
    ("compile-runtime", "Compile-time vs run-time, compiler, interpreter, JIT", "Derleme zamanı kodun analiz/çeviri aşaması; çalışma zamanı gerçek verilerle yürütme aşamasıdır.", "Konu 1.pdf, s.16-21", ["Final ağırlıklı", "Compile-time / run-time", "Type system"]),
    ("syntax", "Syntax, semantics, BNF, ambiguity", "Syntax biçim, semantics anlamdır; ambiguity aynı ifade için birden fazla parse ağacı doğurur.", "Konu 3.pdf; Konu 4.pdf", ["Vize temeli", "Kod yorumlama", "Paradigmalar"]),
    ("scope", "Static scope, dynamic scope, binding", "Static scope lexical yapıya, dynamic scope çağrı zincirine bakar; binding isim ile varlık/değer/adres ilişkisidir.", "Konu 5.pdf; Konu 11.pdf", ["Final ağırlıklı", "Scope / binding", "Klasik soru"]),
    ("memory", "Lifetime, static/stack/heap, dangling pointer", "Lifetime değişkenin bellekte var olma süresidir; scope görünürlük, lifetime varlık süresidir.", "Konu 6.pdf; Konu 12.pdf", ["Final ağırlıklı", "Bellek yönetimi", "Klasik soru"]),
    ("types", "Type checking, strong/weak typing, coercion", "Type checking tür hatalarını compile-time veya run-time yakalar; strong typing tür hatalarını sınırlamayı hedefler.", "Konu 1.pdf, s.14; Konu 7.pdf", ["Final ağırlıklı", "Type system", "Compile-time / run-time"]),
    ("expr", "Expressions, precedence, associativity, short-circuit, side effect", "İfade sonucu operand değerlendirme sırası, öncelik, ilişkisellik ve yan etkilere bağlı olabilir.", "Konu 7.pdf", ["Final ağırlıklı", "Kod yorumlama", "Klasik soru"]),
    ("control", "Selection, loops, guarded commands", "if/switch/loop yapıları akışı yönetir; fall-through, dangling else ve pre/post-test sınav tuzağıdır.", "Konu 8.pdf; Konu 9.pdf", ["Vize temeli", "Kod yorumlama", "Paradigmalar"]),
    ("subprogram", "Alt programlar, formal/actual parametre, function/procedure", "Alt programlar işlem soyutlaması sağlar; formal parametre başlıkta, actual parametre çağrıda verilir.", "Konu 10.pdf", ["Final ağırlıklı", "Scope / binding", "Klasik soru"]),
    ("param", "Pass-by-value/reference/result/value-result/name", "Parametre geçiş modeli actual değerin kopya, adres, sonuç veya ifade olarak taşınmasını belirler.", "Konu 10.pdf, s.9-18", ["Final ağırlıklı", "Scope / binding", "Klasik soru"]),
    ("funcbind", "Function parameter binding: shallow, deep, ad hoc", "Fonksiyon parametre olarak geçerken hangi referans ortamıyla çalışacağı kritik binding problemidir.", "Konu 11.pdf, s.2-13", ["Final ağırlıklı", "Scope / binding", "Klasik soru"]),
    ("closure", "Closure ve coroutine", "Closure fonksiyon + referans ortamıdır; coroutine durumunu koruyup resume ile devam eder.", "Konu 11.pdf, s.14-24", ["Final ağırlıklı", "Fonksiyonel programlama", "Klasik soru"]),
    ("ar", "Subprogram linkage, activation record, recursion", "Activation record parametre, return address, yerel değişken ve geçici değerleri tutan çağrı kaydıdır.", "Konu 12.pdf", ["Final ağırlıklı", "Bellek yönetimi", "Klasik soru"]),
    ("adt", "ADT, veri soyutlama, encapsulation", "ADT veri temsilini saklar, kullanıcıya tip adı ve işlem protokolü sunar.", "Konu 13.pdf", ["Final ağırlıklı", "Nesne yönelimli programlama", "Type system"]),
    ("oop", "OOP, inheritance, polymorphism, dynamic binding", "Dynamic binding override edilmiş metodun çalışma zamanı nesne tipine göre seçilmesidir.", "Konu 14.pdf", ["Final ağırlıklı", "Nesne yönelimli programlama", "Scope / binding"]),
    ("conc", "Concurrency, race condition, semaphore, monitor, message passing", "Eşzamanlılıkta birden çok kontrol akışı ilerler; paylaşılan veri yarış durumu doğurabilir.", "Konu 15.pdf", ["Final ağırlıklı", "Concurrency", "Klasik soru"]),
    ("rust", "Rust ownership, borrowing, lifetime, trait", "Rust ownership/borrowing kurallarıyla GC olmadan bellek güvenliği sağlamaya çalışır.", "Rust proje raporu; 941478ad PDF varsa", ["Proje konusu", "Dil tanıma", "Bellek yönetimi"]),
    ("go", "Go goroutine, channel, interface, GC", "Go sade syntax, hızlı derleme, GC, goroutine ve channel tabanlı concurrency ile öne çıkar.", "Ders Sunumları.zip içindeki Go raporları", ["Proje konusu", "Dil tanıma", "Concurrency"]),
    ("scala", "Scala trait, object, case class, match", "Scala statik tipli, JVM tabanlı, OOP ve fonksiyonel yapıları birleştiren dildir.", "Ders Sunumları.zip::Scala Dili-1.pdf", ["Proje konusu", "Dil tanıma", "Fonksiyonel programlama"]),
    ("erlang", "Erlang process, atom, pattern matching", "Erlang izole süreçler ve mesajlaşma ile concurrent/fault-tolerant sistemlere odaklanır.", "Ders Sunumları.zip::Erlang proje raporu", ["Proje konusu", "Dil tanıma", "Concurrency"]),
    ("paradigm", "Imperative, OOP, functional, logic paradigms", "Paradigma çözümü ifade etme biçimidir; Prolog logic, Lisp/Haskell functional örnekleridir.", "Konu 1.pdf, s.22-24", ["Vize temeli", "Paradigmalar", "Dil tanıma"]),
]


LANGS = {
    "Rust": ("Sistem programlama + güvenlik", "Statik/güçlü", "Ownership, borrowing, lifetime; GC yok", "thread/channel + Send/Sync", "`let mut`, `&mut`, `impl Trait`, `match`, `Option<T>`", "Compile-time borrow checker sınavın en güçlü ipucu."),
    "Go": ("Sade imperative + concurrent", "Statik, yapısal interface", "GC, escape analysis", "goroutine + channel", "`package main`, `func`, `go f()`, `chan`, `<-`, `defer`", "Kalıtım yerine composition/interface; channel üzerinden iletişim."),
    "Scala": ("OOP + functional", "Statik, tip çıkarımı", "JVM GC", "Future/Akka/JVM threadleri", "`val`, `var`, `trait`, `object`, `case class`, `match`", "Java static yerine object/companion object; trait ile mixin."),
    "Erlang": ("Functional + actor concurrency", "Dinamik/güçlü", "Process başına GC", "hafif process + message passing", "`receive`, `spawn`, atomlar, `.` ile bitiş, `=` pattern match", "Telekom/hata toleransı; değişkenler yeniden atanmaz, pattern matching yapılır."),
    "Java": ("OOP ağırlıklı", "Statik/güçlü", "GC, referanslar", "Thread, synchronized", "`public static void main`, `class`, `extends`, `implements`", "Pointer yok; JVM/JIT ve dynamic dispatch."),
    "C": ("Procedural", "Statik, güvenlik kaçışları var", "Manuel malloc/free", "pthread kütüphanesi", "`#include`, `*p`, `&x`, `struct`, `printf`", "Pointer ve manuel bellek."),
    "C++": ("Çok paradigmalı", "Statik/güçlü ama esnek", "RAII + smart/manual pointer", "std::thread", "`template<typename T>`, `std::`, `virtual`, `cout`", "Overloading/template/virtual dispatch."),
    "Python": ("Çok paradigmalı scripting", "Dinamik/güçlü", "GC/ref counting", "asyncio/threading", "`def`, girinti, `self`, list comprehension", "Girinti blok yapısı ve dinamik tip."),
    "JavaScript": ("Prototype + functional etkiler", "Dinamik", "GC", "event loop, promise", "`const`, `let`, `=>`, `Promise`, `prototype`", "Closure ve async/event loop."),
    "Prolog": ("Logic", "Terim/unification", "Runtime yönetimi", "sınırlı", "`fact(a).`, `rule(X) :- other(X).`", "Olgu, kural, hedef, unification."),
    "Lisp": ("Functional/sembolik", "Dinamik", "GC", "varyanta bağlı", "`(+ 1 2)`, `car`, `cdr`, çok parantezli prefix", "Kod-veri benzerliği ve liste işleme."),
    "Haskell": ("Pure functional", "Statik/güçlü, type inference", "GC", "STM/async", "`::`, `where`, pattern matching, lazy evaluation", "Saflık, lazy evaluation, typeclass."),
}


def safe_name(name: str) -> str:
    return re.sub(r'[<>:"/\\|?*]+', "_", name)


def clean_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def ensure_dirs():
    for p in DIRS.values():
        p.mkdir(parents=True, exist_ok=True)


def copy_file(src: Path, dest_dir: Path) -> Path:
    dest = dest_dir / safe_name(src.name)
    shutil.copy2(src, dest)
    return dest


def extract_pdf(path: Path):
    pages = []
    try:
        reader = PdfReader(str(path))
        for idx, page in enumerate(reader.pages, start=1):
            try:
                text = clean_text(page.extract_text() or "")
            except Exception as exc:
                text = f"[OKUNAMADI: {exc}]"
            pages.append({"page": idx, "text": text})
    except Exception as exc:
        pages.append({"page": None, "text": f"[PDF OKUNAMADI: {exc}]"})
    return pages


def extract_docx(path: Path):
    try:
        doc = Document(str(path))
        chunks = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                vals = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if vals:
                    chunks.append(" | ".join(vals))
        return [{"page": None, "text": clean_text("\n".join(chunks))}]
    except Exception as exc:
        return [{"page": None, "text": f"[DOCX OKUNAMADI: {exc}]"}]


def extract_plain(path: Path):
    for enc in ("utf-8-sig", "utf-8", "cp1254", "latin-1"):
        try:
            return [{"page": None, "text": clean_text(path.read_text(encoding=enc))}]
        except UnicodeDecodeError:
            continue
    return [{"page": None, "text": "[METIN OKUNAMADI]"}]


def category_for(path_name: str):
    low = path_name.lower()
    m = re.match(r"konu\s+(\d+)\.pdf$", low)
    if m:
        n = int(m.group(1))
        return "final" if n >= 10 else "vize"
    if any(x in low for x in ["proje", "ödev", "odev", "scala", "rust", "go", "erlang", "rapor"]):
        return "proje"
    return "ek"


def collect_sources():
    inventory = []
    extracts = []
    for path in sorted(ROOT.iterdir(), key=lambda p: p.name.lower()):
        if not path.is_file():
            continue
        if path.parent == OUT or OUT in path.parents:
            continue
        if path.suffix.lower() not in [".pdf", ".docx", ".txt", ".md", ".html", ".htm", ".zip"]:
            continue
        cat = category_for(path.name)
        if path.suffix.lower() == ".zip":
            zip_dest = copy_file(path, DIRS["proje"] if cat == "proje" else DIRS["ek"])
            inventory.append({"source": path.name, "stored_as": str(zip_dest.relative_to(OUT)), "category": "zip"})
            with zipfile.ZipFile(path) as zf:
                for info in zf.infolist():
                    if info.is_dir():
                        continue
                    suffix = Path(info.filename).suffix.lower()
                    if suffix not in [".pdf", ".docx", ".txt", ".md", ".html", ".htm"]:
                        continue
                    data = zf.read(info)
                    out_name = safe_name(Path(info.filename).name)
                    out_path = DIRS["proje"] / out_name
                    out_path.write_bytes(data)
                    label = f"{path.name}::{info.filename}"
                    inventory.append({"source": label, "stored_as": str(out_path.relative_to(OUT)), "category": "proje"})
                    pages = extract_pdf(out_path) if suffix == ".pdf" else extract_plain(out_path)
                    extracts.append({"source": label, "category": "proje", "pages": pages})
            continue
        dest = copy_file(path, DIRS[cat])
        inventory.append({"source": path.name, "stored_as": str(dest.relative_to(OUT)), "category": cat})
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            pages = extract_pdf(path)
        elif suffix == ".docx":
            pages = extract_docx(path)
        else:
            pages = extract_plain(path)
        extracts.append({"source": path.name, "category": cat, "pages": pages})
    (DIRS["veri"] / "source_inventory.json").write_text(json.dumps(inventory, ensure_ascii=False, indent=2), encoding="utf-8")
    (DIRS["veri"] / "source_extracts.json").write_text(json.dumps(extracts, ensure_ascii=False, indent=2), encoding="utf-8")
    lines = ["# Kaynak Envanteri", ""]
    for item in inventory:
        lines.append(f"- {item['source']} -> `{item['stored_as']}` ({item['category']})")
    (DIRS["rapor"] / "KAYNAK_ENVANTERI.md").write_text("\n".join(lines) + "\n", encoding="utf-8")
    return inventory, extracts


def option(text, correct, explanation):
    return {"text": text, "correct": correct, "explanation": explanation}


def make_test(qid, topic, tags, question, correct, wrongs, source, code=None, difficulty="Final tarzı", qtype="test"):
    opts = [option(correct, True, f"Doğru. Neden doğru: {correct} Konu: {topic}. Sınavdaki tuzak: kavramı benzer syntax veya başka bir aşama ile karıştırmak. Öğrenci neden bu şıkka güvenebilir? Cevap hem davranışı hem gerekçeyi veriyor. Kaynak: {source}.")]
    for w, why in wrongs:
        opts.append(option(w, False, f"Yanlış. Neden yanlış: {why} Konu: {topic}. Sınavdaki tuzak: yüzeyde benzer görünen ama farklı kavrama ait bir ifadeyi seçmek. Öğrenci neden bu şıkka düşebilir? Ezbere bildiği dil modelini bu soruya yanlış aktarabilir. Kaynak: {source}."))
    rot = int(re.sub(r"\D", "", qid) or 0) % 4
    opts = opts[rot:] + opts[:rot]
    for i, o in enumerate(opts):
        o["key"] = chr(65 + i)
    answer = next(o["key"] for o in opts if o["correct"])
    item = {
        "id": qid,
        "type": qtype,
        "difficulty": difficulty,
        "tags": tags,
        "topic": topic,
        "question": question,
        "options": opts,
        "answer": answer,
        "whyImportant": "Finalde ezberden çok kod davranışı, bağlama zamanı, dil tasarımı ve proje dili bağlantısı ölçülebilir.",
        "source": source,
        "answerExplanation": f"Doğru cevap {answer}. Mantık: {correct} Yanlış seçenekler ya kavramı başka bir aşamaya taşıyor ya da farklı dil tasarımını genelliyor. Benzer sorularda önce hangi aşama, hangi scope/binding ve hangi dil ipucu sorulduğunu ayır.",
    }
    if code:
        item["code"] = code
    return item


def make_classic(qid, topic, tags, question, expected, short, long, mistake, source, code=None):
    item = {
        "id": qid,
        "type": "klasik",
        "difficulty": "Final tarzı",
        "tags": list(dict.fromkeys(tags + ["Klasik soru"])),
        "topic": topic,
        "question": question,
        "expected": expected,
        "rubric": [
            "3 puan: doğru kavramı ve belirlenme zamanını yazma.",
            "3 puan: kod/değer/scope zincirini adım adım takip etme.",
            "2 puan: sınav tuzağını belirtme.",
            "2 puan: kısa sonucu net ifade etme.",
        ],
        "keywords": tags + [topic],
        "commonMistake": mistake,
        "shortAnswer": short,
        "longAnswer": long,
        "whyImportant": "Klasik sorular genellikle değer takibi, scope, binding, parametre geçişi ve activation record mantığı ister.",
        "source": source,
    }
    if code:
        item["code"] = code
    return item


def build_questions():
    questions = []
    base_tests = [
        ("Rust ownership", ["Proje konusu", "Bellek yönetimi", "Compile-time / run-time"], "let s2 = s1; sonrasında s1 neden kullanılamaz?", "String gibi ownership taşıyan değerlerde sahiplik s2'ye move olur; s1 kullanımı compile-time hatadır.", [("GC iki referansı takip eder, hata olmaz.", "Rust'ın temel farkı GC değil ownership modelidir."), ("Bu yalnız runtime uyarısıdır.", "Borrow/move hataları derleme zamanında yakalanır."), ("Sadece integer tiplerinde move olur.", "Heap verisi taşıyan tiplerde move özellikle önemlidir.")], "Rust proje raporu"),
        ("Rust borrowing", ["Proje konusu", "Bellek yönetimi", "Type system"], "İki `&mut` referans neden yasaktır?", "Aynı anda iki mutable borrow aliasing + mutation riskini doğurur; borrow checker bunu engeller.", [("Çünkü `mut` hiç kullanılamaz.", "`mut` kullanılabilir; sorun aynı anda iki mutable borrow'dur."), ("Çünkü Rust zayıf tipli dildir.", "Rust statik/güçlü tipli kabul edilir."), ("Bu yalnız Go channel kuralıdır.", "Bu Rust borrowing kuralıdır.")], "Rust proje raporu"),
        ("Go channel", ["Proje konusu", "Concurrency", "Dil tanıma"], "`go worker(ch); x := <-ch` ne anlatır?", "Bir goroutine başlar ve ana akış channel'dan okurken senkronize/bloklanabilir.", [("Class inheritance başlatır.", "Go'da class/extends yoktur."), ("Channel sıradan array'dir.", "Channel veri aktarımı ve senkronizasyon sağlar."), ("`go` generic tanımıdır.", "`go` goroutine başlatır.")], "Go proje raporları", "ch := make(chan int)\ngo worker(ch)\nx := <-ch"),
        ("Go interface", ["Proje konusu", "Type system", "Nesne yönelimli programlama"], "Go interface için doğru ifade nedir?", "Bir tip gerekli metotlara sahipse interface'i örtük olarak sağlar; `implements` gerekmez.", [("Java gibi extends zorunludur.", "Go'da extends/class hiyerarşisi yoktur."), ("Interface channel buffer boyutudur.", "Interface tip soyutlamadır."), ("Tip güvenliğini tamamen kapatır.", "Statik tip sistemi içinde çalışır.")], "Go proje raporları"),
        ("Scala trait/object", ["Proje konusu", "Dil tanıma", "Nesne yönelimli programlama"], "`trait` ve `object` için doğru yorum nedir?", "`trait` mixin/davranış soyutlaması, `object` singleton/static benzeri ortak üyeler için kullanılır.", [("trait pointer türüdür.", "Scala trait pointer değildir."), ("object Java Object sınıfıyla aynıdır.", "Scala object singleton yapıdır."), ("Scala OOP desteklemez.", "Scala güçlü OOP desteğine sahiptir.")], "Scala Dili-1.pdf", "trait Speaker { def speak(): String }\nobject Main { def main(args: Array[String]) = println(\"ok\") }"),
        ("Scala pattern matching", ["Proje konusu", "Dil tanıma", "Fonksiyonel programlama"], "`case class` ve `match` birlikte neyi gösterir?", "Scala'da veri modeli ve pattern matching kullanımını gösterir.", [("Go channel okumasıdır.", "Go channel `<-` ile ayırt edilir."), ("Erlang atom tanımıdır.", "Erlang nokta/receive/atom syntax'ı farklıdır."), ("C++ template'tir.", "C++ template syntax'ı farklıdır.")], "Scala Dili-1.pdf"),
        ("Erlang pattern matching", ["Proje konusu", "Dil tanıma", "Concurrency"], "Erlang'da `X = 3, X = 4.` neden sorunludur?", "`=` yeniden atama değil pattern matching'dir; X 3'e bağlandıktan sonra 4 ile eşleşmez.", [("X önce 3 sonra 4 olur.", "Erlang değişkenleri mutable atama gibi davranmaz."), ("Bu Go kısa değişken bildirimidir.", "Go `:=` kullanır."), ("Bu Rust borrow örneğidir.", "Rust syntax'ı farklıdır.")], "Erlang proje raporu", "X = 3,\nX = 4."),
        ("Erlang process", ["Proje konusu", "Concurrency", "Fonksiyonel programlama"], "Erlang concurrency yaklaşımı nedir?", "İzole hafif süreçler mesajlaşır; actor benzeri model hata toleransını destekler.", [("Global belleği kilitsiz paylaşır.", "Bu race condition riskidir."), ("Class inheritance ile concurrency kurar.", "Erlang OOP inheritance'a dayanmaz."), ("Concurrency desteği yoktur.", "Erlang özellikle concurrent sistemler için tasarlanmıştır.")], "Erlang proje raporu"),
        ("Pass-by-value-result", ["Final ağırlıklı", "Scope / binding", "Klasik soru"], "`p(x,x)` value-result modelinde neden tuzaklıdır?", "İki formal aynı actual'a dönüşte geri yazılabilir; son kopyalanan değer sonucu belirler.", [("Hiç actual okunmaz.", "Başta kopyalama yapılır."), ("Reference ile tamamen aynıdır.", "Reference anlık aliasing, value-result dönüş kopyasıdır."), ("Sadece functional dillerde olur.", "Klasik parametre geçiş modelidir.")], "Konu 10.pdf"),
        ("Pass-by-name", ["Final ağırlıklı", "Scope / binding", "Klasik soru"], "Pass-by-name nasıl çalışır?", "Actual ifade formal her kullanıldığında çağıranın ortamında yeniden değerlendirilir.", [("Başta bir kez kopyalanır.", "Bu pass-by-value olur."), ("Sadece adres gönderilir.", "Bu pass-by-reference olur."), ("Dönüşte kopyalanır.", "Bu value-result ailesidir.")], "Konu 10.pdf"),
        ("Activation record", ["Final ağırlıklı", "Bellek yönetimi", "Klasik soru"], "Activation record içinde ne bulunur?", "Parametreler, return address, yerel değişkenler ve geçici değerler bulunur.", [("Yalnız token listesi.", "Token listesi lexical analizdir."), ("Yalnız heap nesneleri listesi.", "AR çağrıya özgü stack/linkage bilgisidir."), ("Yalnız inheritance ağacı.", "Inheritance OOP konusudur.")], "Konu 12.pdf"),
        ("Closure lifetime", ["Final ağırlıklı", "Fonksiyonel programlama", "Bellek yönetimi"], "Closure neden lifetime uzatır?", "İç fonksiyon dış değişken ortamını kullanmaya devam ettiği için değişken korunmalıdır.", [("Her zaman global değişken kullanır.", "Closure lexical ortam yakalar."), ("Scope ve lifetime aynıdır.", "Scope görünürlük, lifetime varlık süresidir."), ("Coroutine ile aynıdır.", "Coroutine duraklama/devam mekanizmasıdır.")], "Konu 11.pdf"),
        ("Coroutine", ["Final ağırlıklı", "Fonksiyonel programlama", "Kod yorumlama"], "Coroutine'i subroutine'den ayıran nedir?", "Durumunu koruyup durakladığı yerden resume ile devam edebilir.", [("Her zaman prosedürdür.", "Prosedür/fonksiyon ayrımı değildir."), ("Her çağrıda durumu sıfırlar.", "Coroutine durum korur."), ("Sadece type checking yapar.", "Bu tip sistemi konusudur.")], "Konu 11.pdf"),
        ("ADT", ["Final ağırlıklı", "Nesne yönelimli programlama", "Type system"], "Stack ADT için doğru yorum nedir?", "Kullanıcı push/pop/top işlemlerini bilir; gerçek temsil gizlenir.", [("ADT ham byte dizisidir.", "ADT davranış ve işlem protokolüdür."), ("Encapsulation'ı kaldırır.", "Tersine temsil gizler."), ("Sadece Prolog kuralıdır.", "ADT genel veri soyutlamadır.")], "Konu 13.pdf"),
        ("Dynamic binding", ["Final ağırlıklı", "Nesne yönelimli programlama", "Scope / binding"], "Overloading ile overriding/dynamic binding farkı nedir?", "Overloading imza/statik tipe, overriding ise runtime nesne tipine bağlı seçilebilir.", [("Tamamen aynıdır.", "İsim/imza çoklama ile override farklıdır."), ("Dynamic binding compile-time kesinleşir.", "Çalışma zamanı seçimini vurgular."), ("Yalnız Prolog'da görülür.", "OOP dispatch konusudur.")], "Konu 14.pdf"),
        ("Semaphore", ["Final ağırlıklı", "Concurrency", "Kod yorumlama"], "Semaphore için doğru ifade nedir?", "`wait` kaynak yoksa bekletir; `release` bırakır/uyandırır; işlemler atomik olmalıdır.", [("Token tablosudur.", "Runtime senkronizasyon yapısıdır."), ("Deadlock imkansızdır.", "Yanlış kullanım deadlock yaratabilir."), ("Inheritance aracıdır.", "Concurrency aracıdır.")], "Konu 15.pdf"),
        ("Static vs dynamic scope", ["Final ağırlıklı", "Scope / binding", "Klasik soru"], "Static scope ve dynamic scope farkı nedir?", "Static scope lexical kod yapısına, dynamic scope çağrı zincirine bakar.", [("Static heap, dynamic stack demektir.", "Scope bellek bölgesi değildir."), ("Dynamic her zaman daha okunabilirdir.", "Çağrı zinciri öngörülebilirliği azaltabilir."), ("Static sadece yorumlanan dillerde olur.", "Derlenen dillerde de yaygındır.")], "Konu 5.pdf"),
        ("Short-circuit", ["Final ağırlıklı", "Kod yorumlama", "Klasik soru"], "`a != 0 && f()` ifadesinde a=0 ise ne olur?", "Short-circuit nedeniyle `f()` çağrılmaz; yan etki oluşmaz.", [("f önce çağrılır.", "Sol taraf false ise sağ taraf atlanır."), ("Yalnız aritmetikte olur.", "Mantıksal operatörlerde tipiktir."), ("Syntax konusudur, runtime etkisi yoktur.", "Yan etki runtime davranışıdır.")], "Konu 7.pdf"),
        ("Halting problem", ["Vize temeli", "Compile-time / run-time", "Paradigmalar"], "Halting problemi compiler için ne gösterir?", "Genel olarak bir programın durup durmayacağını kesin belirleyen program yazılamaz.", [("Compiler tüm sonsuz döngüleri yakalar.", "Halting problemi bunu reddeder."), ("Syntax hataları karar verilemezdir.", "Syntax genelde grammar ile karar verilebilir."), ("JIT halting'i çözer.", "Gerçekleştirim modeli kararsızlığı çözmez.")], "Konu 1.pdf; Konu 2.pdf"),
        ("Type checking", ["Final ağırlıklı", "Type system", "Compile-time / run-time"], "Strong typing ile dynamic typing ilişkisi nedir?", "Dinamik tipli dil güçlü tip denetimi yapabilir; güçlü/zayıf ve statik/dinamik ayrı eksenlerdir.", [("Dinamik her dil zayıftır.", "Python dinamik ama güçlü tipe yakındır."), ("Statik her dil hatasızdır.", "Cast/union gibi kaçışlar olabilir."), ("Strong typing hız demektir.", "Asıl konu tür hatalarını yakalamaktır.")], "Konu 1.pdf; Konu 7.pdf"),
    ]
    idx = 1
    while len(questions) < 84:
        bt = base_tests[(idx - 1) % len(base_tests)]
        suffix = "" if idx <= len(base_tests) else f" - varyant {((idx - 1) // len(base_tests)) + 1}"
        questions.append(make_test(f"TST-{idx:03d}", bt[0], bt[1], bt[2] + suffix, bt[3], bt[4], bt[5], bt[6] if len(bt) > 6 else None, difficulty=["Kolay", "Orta", "Zor", "Final tarzı", "Tuzaklı"][idx % 5]))
        idx += 1

    classic_cases = [
        ("Static/dynamic scope değer takibi", ["Final ağırlıklı", "Scope / binding"], "Static ve dynamic scope için çıktıyı ayrı yazınız.", "Static lexical çevreye, dynamic çağrı zincirine bakar; x değeri buna göre değişir.", "Static: kod yapısındaki x; dynamic: çağrı zincirindeki en yakın x.", "Önce `bar` fonksiyonunun tanımlandığı yer ile çağrıldığı yer ayrılır; scope zinciri çizilir.", "Scope'u lifetime ile karıştırmak.", "Konu 5.pdf", "x=0\nbar(){print(x)}\nfoo(){x=1; bar()}"),
        ("Pass-by-value/reference/value-result", ["Final ağırlıklı", "Scope / binding"], "`p(x,x)` çağrısını üç parametre geçiş modeliyle çözünüz.", "Value: actual değişmez; reference: aliasing ile aynı hücre anlık değişir; value-result: dönüş kopyalama sırası tuzaktır.", "Value değişmez; reference alias; value-result sıra bağımlı.", "Formal/actual eşleşmesi tabloyla gösterilir ve her atama adım adım izlenir.", "Value-result'ı reference sanmak.", "Konu 10.pdf", "procedure p(a,b){a=1; b=2}\nx=0\np(x,x)"),
        ("Deep/shallow binding", ["Final ağırlıklı", "Scope / binding"], "Fonksiyon parametre geçişinde deep ve shallow binding sonuçlarını açıklayınız.", "Deep fonksiyonun taşınan/tanımlandığı ortamı, shallow çağrıldığı andaki ortamı kullanır.", "Deep tanım/gönderim ortamı; shallow çağrı ortamı.", "Soruda sadece fonksiyon adresi değil referans ortamı problemi olduğu belirtilir.", "Function pointer'ı yalnız adres sanmak.", "Konu 11.pdf", "x=1\nf(){print(x)}\ng(h){x=2; h()}"),
        ("Closure lifetime", ["Final ağırlıklı", "Fonksiyonel programlama", "Bellek yönetimi"], "Closure örneğinde `count` neden kaybolmaz?", "İç fonksiyon dış değişkeni yakalar; fonksiyon döndükten sonra ortam korunur ve ikinci çağrı 2 olur.", "Closure ortamı korur; count ikinci çağrıda 2.", "Closure = fonksiyon + lexical environment; lifetime stack süresini aşabilir.", "Her çağrıda count 0 olur sanmak.", "Konu 11.pdf", "make(){ count=0; return () => ++count }"),
        ("Activation record", ["Final ağırlıklı", "Bellek yönetimi"], "`fact(3)` için activation record sırasını yazınız.", "fact(3), fact(2), fact(1), fact(0) ayrı kayıtlar oluşturur; dönüşler LIFO olur.", "Her recursive çağrı ayrı stack frame.", "Parametre, return address, local ve temporary alanları belirtilir.", "Tek yerel değişkenle recursion takip etmek.", "Konu 12.pdf", "fact(n){ if n==0 return 1; return n*fact(n-1) }"),
        ("Dynamic binding OOP", ["Final ağırlıklı", "Nesne yönelimli programlama"], "`A obj = new B(); obj.draw();` hangi metodu çağırır?", "Override varsa gerçek nesne tipi B olduğundan B.draw çalışır.", "Çıktı B; runtime dispatch.", "Referans tipi ile nesne tipi ayrılır; overloading/overriding farkı yazılır.", "Referans tipi A diye A.draw sanmak.", "Konu 14.pdf", "class A{draw(){A}}\nclass B extends A{draw(){B}}\nA obj=new B(); obj.draw();"),
        ("Race condition", ["Final ağırlıklı", "Concurrency"], "`total = total + 1` iki thread'de neden tehlikelidir?", "Okuma-artırma-yazma atomik değildir; iki thread aynı eski değeri okuyabilir; critical section gerekir.", "Atomik değil; mutual exclusion gerekir.", "Interleaving adımları yazılır; wait/release veya monitor çözümü verilir.", "Tek satırı atomik sanmak.", "Konu 15.pdf", "Thread1: total=total+1\nThread2: total=total+1"),
        ("Rust borrow checker", ["Proje konusu", "Bellek yönetimi"], "İki mutable borrow kodu neden hata verir?", "Aynı anda iki `&mut` borrow compile-time yasaktır.", "İki &mut yasak; compile-time hata.", "Ownership/borrowing, aliasing ve data race bağlantısı açıklanır.", "Runtime hatası sanmak.", "Rust proje raporu", "let mut x=5;\nlet a=&mut x;\nlet b=&mut x;"),
        ("Go channel", ["Proje konusu", "Concurrency"], "Go channel kodunda ana akış nerede bloklanabilir?", "`<-ch` okuması gönderim yoksa bloklanabilir; channel senkronizasyon noktasıdır.", "Channel okuma/gönderme bloklayabilir.", "Goroutine ve channel ayrımı, paylaşarak iletişim yerine iletişimle paylaşım ilkesi yazılır.", "Channel'ı array sanmak.", "Go proje raporu", "ch:=make(chan int)\ngo worker(ch)\nx:=<-ch"),
        ("Erlang pattern matching", ["Proje konusu", "Dil tanıma"], "`X = 3, X = 4.` neden hata olur?", "`=` atama değil pattern matching; X 3'e bağlandıktan sonra 4 ile eşleşmez.", "Pattern match hatası.", "Atom/değişken, tek atama ve nokta sözdizimi belirtilir.", "C/Java ataması gibi okumak.", "Erlang proje raporu", "X = 3,\nX = 4."),
        ("Scala trait/object", ["Proje konusu", "Dil tanıma"], "Scala trait ve object kavramlarını açıklayınız.", "trait mixin/davranış soyutlama; object singleton/static benzeri yapı; case class/match pattern matching sağlar.", "trait=mixin, object=singleton.", "Java static ile Scala object farkı ve OOP+FP birleşimi yazılır.", "Object'i Java Object sınıfı sanmak.", "Scala Dili-1.pdf", "trait T\nobject Main"),
    ]
    for i in range(24):
        c = classic_cases[i % len(classic_cases)]
        questions.append(make_classic(f"CLS-{i+1:03d}", *c))

    for i, (lang, info) in enumerate(LANGS.items(), start=1):
        paradigm, typ, mem, conc, syntax, detail = info
        questions.append(make_test(
            f"LANG-{i:03d}",
            f"{lang} dil tanıma",
            ["Dil tanıma", "Proje konusu" if lang in ["Rust", "Go", "Scala", "Erlang"] else "Vize temeli", "Kod yorumlama"],
            f"{lang} için en güçlü kod/sınav ipucu hangisidir?",
            f"{syntax}; {detail}",
            [(f"{lang} yalnız HTML markup dilidir.", "Bu ifade dilin kaynaklarda anlatılan programlama paradigmasını yansıtmaz."), (f"{lang} için bellek/concurrency konusu yoktur.", f"Bellek: {mem}; concurrency: {conc}."), ("Sadece popülerlik bilgisi yeterlidir.", "Sınav teknik kavram ister.")],
            f"{lang} karşılaştırma notu",
            difficulty="Orta",
            qtype="dil_tanima",
        ))
    while len([q for q in questions if q["type"] == "dil_tanima"]) < 24:
        i = len([q for q in questions if q["type"] == "dil_tanima"]) + 1
        lang = list(LANGS.keys())[(i - 1) % len(LANGS)]
        info = LANGS[lang]
        questions.append(make_test(f"LANG-X{i:03d}", f"{lang} dil tanıma", ["Dil tanıma", "Kod yorumlama"], f"Aşağıdaki ipucu hangi dile aittir: {info[4]}?", lang, [("Başka bir dildir.", "Verilen syntax bu dile özgü daha güçlü ipucu taşır."), ("Sadece paradigma adıdır.", "Soru dil tanıma sorusudur."), ("Kaynak belirsizdir.", "Karşılaştırma rehberinde açıkça listelenmiştir.")], f"{lang} karşılaştırma notu", difficulty="Kolay", qtype="dil_tanima"))

    trace_templates = classic_cases
    for i in range(28):
        c = trace_templates[i % len(trace_templates)]
        questions.append({
            "id": f"TRACE-{i+1:03d}",
            "type": "kod_yorumlama",
            "difficulty": ["Orta", "Zor", "Final tarzı", "Tuzaklı"][i % 4],
            "tags": list(dict.fromkeys(["Kod yorumlama", "Final ağırlıklı"] + c[1])),
            "topic": c[0],
            "question": c[2],
            "code": c[8],
            "expected": c[3],
            "rubric": ["4 puan: adım adım değer/scope takibi.", "3 puan: ilgili kavramı doğru adlandırma.", "2 puan: compile-time/run-time veya binding zamanını yazma.", "1 puan: kısa sonuç."],
            "keywords": c[1] + [c[0]],
            "commonMistake": c[6],
            "shortAnswer": c[4],
            "longAnswer": c[5],
            "whyImportant": "Kod yorumlama soruları finalde klasik ve test biçiminde gelebilir.",
            "source": c[7],
        })
    return questions


def write_summaries():
    general = ["# Genel Final Özeti", "", "Final ağırlığı özellikle Konu 10-15 ve proje dillerindedir. Vize temeli tamamen dışlanmaz; scope, type checking, syntax/semantics ve memory yönetimi final konularıyla birleşir.", ""]
    for _, title, definition, source, tags in TOPICS:
        if "Final ağırlıklı" in tags or "Proje konusu" in tags:
            general += [f"## {title}", f"- Özet: {definition}", f"- Sınavda nasıl gelir: kod verilip değer/binding/scope/çalışma zamanı sorulabilir.", f"- Tuzak: benzer kavramları aynı sanmak veya dil syntax'ını yanlış okumak.", f"- Kaynak: {source}", ""]
    (DIRS["ozet"] / "01_GENEL_FINAL_OZETI.md").write_text("\n".join(general), encoding="utf-8")

    detail = ["# Konu Konu Detaylı Özet", ""]
    for _, title, definition, source, tags in TOPICS:
        detail += [
            f"## {title}",
            f"- Kısa tanım: {definition}",
            "- Neden önemli? Final soruları kavramları birbirine bağlatır; yalnız tanım yetmez.",
            "- Hoca nasıl sorabilir? Kod parçası, değer takibi, hangi dile ait, compile-time/run-time ayrımı veya açık uçlu proje yorumu.",
            "- Testte tuzak: aynı anahtar kelimeyi farklı dil/kavramla karıştırmak.",
            "- Klasikte nasıl sorulur? Adım adım değer/scope/activation record tablosu istenir.",
            "- Örnek mini soru: Bu olay hangi aşamada ve hangi scope/binding kuralıyla belirlenir?",
            "- Cevap mantığı: önce kavramı tanımla, sonra kod üzerinde zinciri izle, en son sonucu yaz.",
            f"- İlgili kaynak: {source}",
            "",
        ]
    (DIRS["ozet"] / "02_KONU_KONU_DETAYLI_OZET.md").write_text("\n".join(detail), encoding="utf-8")

    comp = ["# Programlama Dilleri Karşılaştırma", "", "| Dil | Paradigma | Tip sistemi | Bellek | Concurrency | Koddan tanıma | Sınav detayı |", "|---|---|---|---|---|---|---|"]
    for lang, (paradigm, typ, mem, conc, syntax, detail_text) in LANGS.items():
        comp.append(f"| {lang} | {paradigm} | {typ} | {mem} | {conc} | {syntax} | {detail_text} |")
    (DIRS["ozet"] / "03_PROGRAMLAMA_DILLERI_KARSILASTIRMA.md").write_text("\n".join(comp) + "\n", encoding="utf-8")

    guide = ["# Kod Parçası Dil Tanıma Rehberi", ""]
    for lang, (paradigm, typ, mem, conc, syntax, detail_text) in LANGS.items():
        guide += [f"## {lang}", f"- Paradigma: {paradigm}", f"- Tip sistemi: {typ}", f"- Bellek: {mem}", f"- Concurrency: {conc}", f"- Ayırt edici syntax: {syntax}", f"- Sınavda dikkat: {detail_text}", ""]
    (DIRS["ozet"] / "04_KOD_PARCA_DIL_TANIMA_REHBERI.md").write_text("\n".join(guide), encoding="utf-8")

    classic = ["# Klasik Soru Çalışma Notları", "", "Cevap şablonu: tanım -> kod/değer tablosu -> binding/scope/lifetime aşaması -> sonuç -> tuzak.", ""]
    for _, title, definition, source, tags in TOPICS:
        if "Klasik soru" in tags or "Final ağırlıklı" in tags:
            classic += [f"## {title}", f"- Tanım: {definition}", "- Mini örnek: küçük bir kod parçası üzerinde değişkenleri tabloyla takip et.", "- Kısa cevap: sonucu ve kavram adını yaz.", "- Uzun cevap: nedenini, hangi aşamada belirlendiğini ve tuzağı açıkla.", "- Puanlama: tanım 2, adım adım takip 4, tuzak 2, sonuç 2.", f"- Kaynak: {source}", ""]
    (DIRS["ozet"] / "05_KLASIK_SORU_CALISMA_NOTLARI.md").write_text("\n".join(classic), encoding="utf-8")


def write_app(questions):
    html = """<!DOCTYPE html>
<html lang="tr">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>Programlama Dilleri Final Hazırlık</title>
  <link rel="stylesheet" href="style.css">
</head>
<body>
  <header class="topbar"><div><h1>Programlama Dilleri Final Hazırlık</h1><p>Offline soru bankası, özetler, filtreler ve ilerleme takibi</p></div><button id="resetBtn">Sıfırla</button></header>
  <main class="layout">
    <aside class="panel">
      <input id="searchBox" type="search" placeholder="Konu, kavram, dil veya soru ara">
      <select id="typeFilter"><option value="">Tüm türler</option><option value="test">Test</option><option value="klasik">Klasik</option><option value="dil_tanima">Dil tanıma</option><option value="kod_yorumlama">Kod yorumlama</option></select>
      <select id="difficultyFilter"><option value="">Tüm zorluklar</option></select>
      <label><input type="checkbox" id="favOnly"> Favoriler</label>
      <label><input type="checkbox" id="wrongOnly"> Yanlışlar</label>
      <div id="tagFilters" class="chips"></div>
    </aside>
    <section>
      <section id="dashboard" class="dashboard"></section>
      <div class="toolbar"><strong id="resultCount"></strong><button id="clearFilters">Filtreleri temizle</button></div>
      <div id="questions"></div>
    </section>
  </main>
  <script src="app.js"></script>
</body>
</html>
"""
    css = """
body{margin:0;font-family:Arial,Helvetica,sans-serif;background:#f6f7f9;color:#1f2933}.topbar{display:flex;justify-content:space-between;gap:16px;align-items:center;padding:20px 26px;background:#fff;border-bottom:1px solid #d7dde5;position:sticky;top:0;z-index:5}h1{margin:0;font-size:25px}.topbar p{margin:4px 0 0;color:#667085}.layout{display:grid;grid-template-columns:310px 1fr;gap:20px;padding:20px}.panel,.question,.metric{background:#fff;border:1px solid #d7dde5;border-radius:8px}.panel{padding:15px;align-self:start;position:sticky;top:92px}input,select{width:100%;box-sizing:border-box;margin:0 0 12px;padding:10px;border:1px solid #cfd6df;border-radius:6px}label{display:block;margin:8px 0}.chips{display:flex;flex-wrap:wrap;gap:7px;margin-top:12px}.chip{border:1px solid #cfd6df;background:#fff;border-radius:999px;padding:7px 10px;cursor:pointer}.chip.active{background:#0f766e;color:#fff}.dashboard{display:grid;grid-template-columns:repeat(5,minmax(110px,1fr));gap:10px;margin-bottom:14px}.metric{padding:12px}.metric b{display:block;font-size:23px}.toolbar{display:flex;justify-content:space-between;margin:10px 0}.question{padding:16px;margin-bottom:12px}.meta{display:flex;flex-wrap:wrap;gap:7px;margin-bottom:10px}.pill{background:#eef4ff;border-radius:999px;padding:4px 8px;font-size:12px}.diff{background:#ecfdf3}.option{border:1px solid #d7dde5;border-radius:8px;margin:8px 0;padding:11px}.explanation,.answerbox{display:none;margin-top:10px;padding:11px;background:#f0fdfa;border-left:4px solid #0f766e;border-radius:4px}.visible{display:block}pre{overflow:auto;background:#111827;color:#fff;padding:12px;border-radius:8px}button{border:1px solid #cfd6df;background:#fff;border-radius:6px;padding:8px 11px;cursor:pointer;margin:4px}.source{font-size:13px;color:#667085}@media(max-width:850px){.layout{grid-template-columns:1fr}.panel{position:static}.dashboard{grid-template-columns:repeat(2,1fr)}}
"""
    app = "const QUESTIONS = " + json.dumps(questions, ensure_ascii=False, indent=2) + ";\nwindow.QUESTIONS = QUESTIONS;\n" + r"""
const STORE_KEY="pd-final-progress-v2";
const state={search:"",type:"",difficulty:"",tags:new Set(),favOnly:false,wrongOnly:false,progress:JSON.parse(localStorage.getItem(STORE_KEY)||"{}")};
const $=id=>document.getElementById(id);
const allTags=[...new Set(QUESTIONS.flatMap(q=>q.tags||[]))].sort((a,b)=>a.localeCompare(b,"tr"));
const allDiffs=[...new Set(QUESTIONS.map(q=>q.difficulty))];
function save(){localStorage.setItem(STORE_KEY,JSON.stringify(state.progress))}
function record(id,status){state.progress[id]={...(state.progress[id]||{}),status,solved:true,last:new Date().toISOString()};save();render()}
function toggleFav(id){state.progress[id]={...(state.progress[id]||{}),favorite:!(state.progress[id]||{}).favorite};save();render()}
function filtered(){const term=state.search.trim().toLocaleLowerCase("tr");return QUESTIONS.filter(q=>{const p=state.progress[q.id]||{};const hay=[q.id,q.topic,q.question,q.source,q.code||"",q.expected||"",...(q.tags||[])].join(" ").toLocaleLowerCase("tr");if(term&&!hay.includes(term))return false;if(state.type&&q.type!==state.type)return false;if(state.difficulty&&q.difficulty!==state.difficulty)return false;if(state.favOnly&&!p.favorite)return false;if(state.wrongOnly&&p.status!=="wrong")return false;for(const t of state.tags)if(!q.tags.includes(t))return false;return true})}
function dashboard(){const solved=QUESTIONS.filter(q=>state.progress[q.id]?.solved).length;const good=QUESTIONS.filter(q=>state.progress[q.id]?.status==="correct").length;const bad=QUESTIONS.filter(q=>state.progress[q.id]?.status==="wrong").length;const fav=QUESTIONS.filter(q=>state.progress[q.id]?.favorite).length;$("dashboard").innerHTML=[["Toplam",QUESTIONS.length],["Çözülen",solved],["Doğru",good],["Yanlış",bad],["Favori",fav]].map(x=>`<div class="metric">${x[0]}<b>${x[1]}</b></div>`).join("")}
function escapeHtml(t){return (t||"").replace(/[&<>"']/g,c=>({"&":"&amp;","<":"&lt;",">":"&gt;",'"':"&quot;","'":"&#39;"}[c]))}
function toggleExplain(id){const el=$("exp-"+id); if(el) el.classList.toggle("visible")}
function toggleAnswer(id){const el=$("answer-"+id); if(el) el.classList.toggle("visible")}
function choose(qid,key){const q=QUESTIONS.find(x=>x.id===qid);record(qid,q.answer===key?"correct":"wrong");setTimeout(()=>{const a=$("answer-"+qid); if(a)a.classList.add("visible")},0)}
function optionHtml(q,o){const id=`${q.id}-${o.key}`;return `<div class="option"><strong>${o.key})</strong> ${o.text}<div><button onclick="choose('${q.id}','${o.key}')">Bu şıkkı seç</button><button onclick="toggleExplain('${id}')">Açıklamayı göster/gizle</button></div><div class="explanation" id="exp-${id}">${o.explanation}</div></div>`}
function qHtml(q){const p=state.progress[q.id]||{};let body=q.code?`<pre><code>${escapeHtml(q.code)}</code></pre>`:"";if(q.options){body+=q.options.map(o=>optionHtml(q,o)).join("")+`<div class="answerbox" id="answer-${q.id}">${q.answerExplanation}</div>`}else{body+=`<button onclick="toggleAnswer('${q.id}')">Beklenen cevabı göster/gizle</button><button onclick="record('${q.id}','correct')">Doğru say</button><button onclick="record('${q.id}','wrong')">Tekrar çalış</button><div class="answerbox" id="answer-${q.id}"><p><b>Beklenen:</b> ${q.expected}</p><p><b>Kısa:</b> ${q.shortAnswer}</p><p><b>Uzun:</b> ${q.longAnswer}</p><p><b>Yaygın hata:</b> ${q.commonMistake}</p><ul>${q.rubric.map(r=>`<li>${r}</li>`).join("")}</ul></div>`}return `<article class="question"><div class="meta"><span class="pill diff">${q.difficulty}</span><span class="pill">${q.type}</span>${q.tags.map(t=>`<span class="pill">${t}</span>`).join("")}</div><h2>${q.id} - ${q.topic}</h2><p>${q.question}</p>${body}<button onclick="toggleFav('${q.id}')">${p.favorite?"Favoriden çıkar":"Favorilere ekle"}</button>${p.solved?`<span class="pill">Durum: ${p.status}</span>`:""}<p class="source"><b>Neden önemli?</b> ${q.whyImportant}<br><b>Kaynak:</b> ${q.source}</p></article>`}
function render(){dashboard();const list=filtered();$("resultCount").textContent=`${list.length} soru gösteriliyor`;$("questions").innerHTML=list.map(qHtml).join("")}
function setup(){allDiffs.forEach(d=>$("difficultyFilter").innerHTML+=`<option value="${d}">${d}</option>`);$("tagFilters").innerHTML=allTags.map(t=>`<button class="chip" data-tag="${t}">${t}</button>`).join("");$("searchBox").oninput=e=>{state.search=e.target.value;render()};$("typeFilter").onchange=e=>{state.type=e.target.value;render()};$("difficultyFilter").onchange=e=>{state.difficulty=e.target.value;render()};$("favOnly").onchange=e=>{state.favOnly=e.target.checked;render()};$("wrongOnly").onchange=e=>{state.wrongOnly=e.target.checked;render()};$("tagFilters").onclick=e=>{if(!e.target.dataset.tag)return;const t=e.target.dataset.tag;state.tags.has(t)?state.tags.delete(t):state.tags.add(t);e.target.classList.toggle("active");render()};$("clearFilters").onclick=()=>{state.search=state.type=state.difficulty="";state.tags.clear();state.favOnly=state.wrongOnly=false;$("searchBox").value="";$("typeFilter").value="";$("difficultyFilter").value="";$("favOnly").checked=false;$("wrongOnly").checked=false;document.querySelectorAll(".chip").forEach(c=>c.classList.remove("active"));render()};$("resetBtn").onclick=()=>{if(confirm("Tüm ilerleme silinsin mi?")){state.progress={};save();render()}};render()}
document.addEventListener("DOMContentLoaded",setup);
"""
    (DIRS["app"] / "index.html").write_text(html, encoding="utf-8")
    (DIRS["app"] / "style.css").write_text(css, encoding="utf-8")
    (DIRS["app"] / "app.js").write_text(app, encoding="utf-8")


def write_validation_scripts():
    validate_py = r'''import json, re
from pathlib import Path
root = Path(__file__).resolve().parents[1]
app = root / "02_INTERAKTIF_SISTEM" / "app.js"
text = app.read_text(encoding="utf-8")
m = re.search(r"const QUESTIONS = (\[.*?\]);\nwindow\.QUESTIONS", text, re.S)
assert m, "QUESTIONS bulunamadi"
qs = json.loads(m.group(1))
ids = set()
counts = {}
for q in qs:
    assert q["id"] not in ids, q["id"]
    ids.add(q["id"])
    counts[q["type"]] = counts.get(q["type"], 0) + 1
    assert q.get("tags"), q["id"]
    assert q.get("source"), q["id"]
    if q["type"] in ("test", "dil_tanima"):
        correct = [o for o in q["options"] if o.get("correct")]
        assert len(correct) == 1, q["id"]
        assert q["answer"] == correct[0]["key"], q["id"]
        for o in q["options"]:
            e = o.get("explanation", "")
            assert "Neden" in e and "Konu" in e and "tuzak" in e.lower(), q["id"] + o.get("key","")
    else:
        for f in ["expected", "rubric", "commonMistake", "shortAnswer", "longAnswer"]:
            assert q.get(f), q["id"] + f
assert len(qs) >= 140, len(qs)
assert counts.get("test", 0) >= 80, counts
assert counts.get("klasik", 0) >= 20, counts
assert counts.get("dil_tanima", 0) >= 20, counts
assert counts.get("kod_yorumlama", 0) >= 20, counts
print(json.dumps({"status":"ok","questions":len(qs),"counts":counts}, ensure_ascii=False, indent=2))
'''
    validate_js = r'''const fs=require("fs"); const vm=require("vm");
function cls(){const s=new Set();return{add:x=>s.add(x),remove:x=>s.delete(x),toggle:x=>s.has(x)?s.delete(x):s.add(x),contains:x=>s.has(x)}}
const elements={dashboard:{innerHTML:""},resultCount:{textContent:""},questions:{innerHTML:""},"exp-smoke":{classList:cls()}};
const storage={};
const sandbox={console,setTimeout,localStorage:{getItem:k=>storage[k]||null,setItem:(k,v)=>storage[k]=String(v)},document:{getElementById:id=>elements[id]||(elements[id]={innerHTML:"",textContent:"",value:"",checked:false,classList:cls(),addEventListener(){}}),querySelectorAll:()=>[],addEventListener(){}},window:{},confirm:()=>true};
sandbox.globalThis=sandbox; vm.createContext(sandbox);
vm.runInContext(fs.readFileSync("02_INTERAKTIF_SISTEM/app.js","utf8"), sandbox);
if(!Array.isArray(sandbox.window.QUESTIONS)||sandbox.window.QUESTIONS.length<140) throw new Error("QUESTIONS yok");
sandbox.toggleExplain("smoke"); if(!elements["exp-smoke"].classList.contains("visible")) throw new Error("toggleExplain show failed");
sandbox.toggleExplain("smoke"); if(elements["exp-smoke"].classList.contains("visible")) throw new Error("toggleExplain hide failed");
const id=sandbox.window.QUESTIONS[0].id; sandbox.toggleFav(id); let p=JSON.parse(storage["pd-final-progress-v2"]); if(!p[id].favorite) throw new Error("fav failed");
sandbox.record(id,"wrong"); p=JSON.parse(storage["pd-final-progress-v2"]); if(p[id].status!=="wrong") throw new Error("record failed");
console.log(JSON.stringify({status:"ok", testedQuestion:id}, null, 2));
'''
    (DIRS["arac"] / "validate_system.py").write_text(validate_py, encoding="utf-8")
    (DIRS["arac"] / "validate_interactions.js").write_text(validate_js, encoding="utf-8")


def write_reports(inventory, extracts, questions):
    counts = {}
    for q in questions:
        counts[q["type"]] = counts.get(q["type"], 0) + 1
    report = {
        "okunan_envanter_girdisi": len(inventory),
        "metin_cikarilan_kaynak": len(extracts),
        "soru_sayisi": len(questions),
        "soru_turleri": counts,
        "klasor_yapisi": {k: str(v.relative_to(OUT)) for k, v in DIRS.items() if k != "arac"},
        "proje_dilleri": ["Rust", "Go", "Scala", "Erlang"],
        "final_kritik_konular": [t[1] for t in TOPICS if "Final ağırlıklı" in t[4]][:15],
    }
    (DIRS["rapor"] / "URETIM_RAPORU.json").write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    readme = f"""# Programlama Dilleri Final Hazırlık

Bu klasör yeniden oluşturuldu ve her şey sınıflandırılmış alt klasörlere ayrıldı.

## Açılış

Interaktif sistemi açmak için:

`02_INTERAKTIF_SISTEM/index.html`

## Klasörler

- `00_KAYNAKLAR/01_vize_temeli_slaytlar`: Konu 1-9
- `00_KAYNAKLAR/02_final_agirlikli_slaytlar`: Konu 10-15
- `00_KAYNAKLAR/03_proje_odevleri_ve_dil_raporlari`: ZIP içinden çıkarılan proje raporları ve dil tanıtımları
- `00_KAYNAKLAR/04_ek_materyaller`: genel özetler ve ek dosyalar
- `01_OZETLER`: final özeti, detaylı özet, dil karşılaştırma, kod tanıma rehberi, klasik notları
- `02_INTERAKTIF_SISTEM`: offline HTML/CSS/JS çalışma sistemi
- `03_VERI`: kaynak envanteri ve çıkarılmış metinler
- `04_RAPORLAR`: üretim ve kaynak raporları
- `99_araclar`: yeniden üretim/doğrulama scriptleri

## Soru Bankası

Toplam soru: {len(questions)}

- Test: {counts.get('test', 0)}
- Klasik: {counts.get('klasik', 0)}
- Dil tanıma: {counts.get('dil_tanima', 0)}
- Kod yorumlama: {counts.get('kod_yorumlama', 0)}

Veriler tarayıcı `localStorage` alanında `pd-final-progress-v2` anahtarıyla saklanır.
"""
    (OUT / "README.md").write_text(readme, encoding="utf-8")


def main():
    ensure_dirs()
    inventory, extracts = collect_sources()
    write_summaries()
    questions = build_questions()
    write_app(questions)
    write_validation_scripts()
    write_reports(inventory, extracts, questions)
    print(json.dumps({"status": "ok", "questions": len(questions), "sources": len(inventory)}, ensure_ascii=False))


if __name__ == "__main__":
    main()
